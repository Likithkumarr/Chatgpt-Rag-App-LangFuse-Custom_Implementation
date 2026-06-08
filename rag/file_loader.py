import os
import uuid
from PIL import Image
import pytesseract

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
 
 
def load_files(files):
    docs = []
 
    # create temp folder
    os.makedirs("temp", exist_ok=True) 
 
    for f in files:
        # 🔥 unique filename (fix overwrite issue)
        unique_name = f"{uuid.uuid4()}_{f.name}"
        path = f"temp/{unique_name}"
 
        # save file
        with open(path, "wb") as t:
            t.write(f.getbuffer())
 
        try:
            # ✅ PDF
            if f.name.endswith(".pdf"):
                docs.extend(PyPDFLoader(path).load())
 
            # ✅ TXT
            elif f.name.endswith(".txt"):
                docs.extend(TextLoader(path).load())
 
            # ✅ DOCX
            elif f.name.endswith(".docx"):
                doc = DocxDocument(path)
                text = "\n".join([para.text for para in doc.paragraphs])
                docs.append(Document(page_content=text, metadata={"source": f.name}))
 
            # ✅ Images (OCR)
            elif f.name.lower().endswith((".png", ".jpg", ".jpeg")):
                text = pytesseract.image_to_string(Image.open(path))
                docs.append(Document(page_content=text, metadata={"source": f.name}))
 
            else:
                print(f"Unsupported file type: {f.name}")
                continue
 
        except Exception as e:
            print(f"Error processing {f.name}: {e}")
 
    return docs